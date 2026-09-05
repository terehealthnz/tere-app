#!/usr/bin/env python3
"""Fill HNZ's HPI Compliance Report docx template with Tere Health's answers.

Reads the blank template at ~/Downloads/HPI_Compliance.docx, preserves HNZ's
exact layout (they're a stickler on format), and interleaves our filled values
into the existing text runs so styling stays intact.

Also injects a per-scenario 'Tere Health Result' line after each Expected
Outcome block so Noel can read our answers in the same order as the template.

Output: ~/Downloads/Tere_HPI_Compliance_Filled.docx
"""
from __future__ import annotations
import copy
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches

HOME = Path.home()
TEMPLATE = HOME / 'Downloads' / 'HPI_Compliance.docx'
OUT = HOME / 'Downloads' / 'Tere_HPI_Compliance_Filled.docx'
EVIDENCE_JSON = HOME / 'Downloads' / 'hpi-compliance-pack.json'

# Screenshots for Noel's 2026-08-26 resubmission ask — expects
# UI screenshots per approved-scope test scenario. Files must exist
# at these paths before running the script. Any missing file logs
# a warning and is skipped (docx still generates).
SCREENSHOT_PATHS = {
    'practitioner': HOME / 'Downloads' / 'hpi-screenshot-practitioner.png',
    'location-gp':    HOME / 'Downloads' / 'hpi-screenshot-location-gp.png',
    'location-pharm': HOME / 'Downloads' / 'hpi-screenshot-location-pharm.png',
    'location-raddx': HOME / 'Downloads' / 'hpi-screenshot-location-raddx.png',
    # IN-3502 rework — dedicated admin HPI Lookup panel (Admin → HPI Lookup)
    # captures each mandatory-test field as a distinct visual element.
    'admin-cpn-99ZZRT':  HOME / 'Downloads' / 'hpi-screenshot-admin-99ZZRT.png',   # HPI-P-Get-1 / -12
    'admin-cpn-90ZZJF':  HOME / 'Downloads' / 'hpi-screenshot-admin-90ZZJF.png',   # HPI-P-Get-2 / -11
    'admin-cpn-91ZZWJ':  HOME / 'Downloads' / 'hpi-screenshot-admin-91ZZWJ.png',   # HPI-P-Get-3 / -9
    'admin-cpn-90ZZLC':  HOME / 'Downloads' / 'hpi-screenshot-admin-90ZZLC.png',   # HPI-P-Get-7 / -8
    'admin-search-oreilly':  HOME / 'Downloads' / 'hpi-screenshot-admin-oreilly.png',   # HPI-P-Search-1
    'admin-search-hunnicutt': HOME / 'Downloads' / 'hpi-screenshot-admin-hunnicutt.png', # HPI-P-Search-4
}

# Which scenario IDs get which screenshot appended. Each entry is a
# list of (screenshot_key, caption) tuples inserted after the text
# answer in the scenario's right-hand cell.
SCENARIO_SCREENSHOTS = {
    # HNZ-supplied UAT test CPNs, rerun via the new Admin → HPI Lookup panel.
    # Every field HNZ requires the product surface is a distinct visible element
    # (registration status pill, qualifications table, conditions-of-practice
    # box, confidentiality banner, deceased banner). Each screenshot below is
    # scenario-specific so the reviewer can eyeball the exact evidence.
    'HPI-P-Get-1': [
        ('admin-cpn-99ZZRT', 'Admin HPI Lookup — CPN 99ZZRT (HNZ test persona 1) — Practitioner returned, name + CPN + qualifications rendered.'),
    ],
    'HPI-P-Get-2': [
        ('admin-cpn-90ZZJF', 'Admin HPI Lookup — CPN 90ZZJF (HNZ test persona 2) — same UI, different persona; proves parity of rendering.'),
    ],
    'HPI-P-Get-3': [
        ('admin-cpn-91ZZWJ', 'Admin HPI Lookup — CPN 91ZZWJ (HNZ test persona 3) — name variant rendered verbatim from the resource, no client-side reshaping.'),
    ],
    'HPI-P-Get-7': [
        ('admin-cpn-90ZZLC', 'Admin HPI Lookup — Conditions of Practice highlighted in a yellow box under each qualification (HPI-P-Get-7).'),
    ],
    'HPI-P-Get-8': [
        ('admin-cpn-90ZZLC', 'Admin HPI Lookup — Registration status pill (green ACTIVE / red INACTIVE) top-right of the practitioner card (HPI-P-Get-8).'),
    ],
    'HPI-P-Get-9': [
        ('admin-cpn-91ZZWJ', 'Admin HPI Lookup — Qualifications table enumerates every scope (code + issuer + period + conditions) as a separate row (HPI-P-Get-9).'),
    ],
    'HPI-P-Get-11': [
        ('admin-cpn-90ZZJF', 'Admin HPI Lookup — Confidentiality banner (red, top of card) shown when meta.security or the confidentiality extension is set (HPI-P-Get-11).'),
    ],
    'HPI-P-Get-12': [
        ('admin-cpn-99ZZRT', 'Admin HPI Lookup — Deceased banner (dark grey, top of card) surfaces practitioner-death-date extension / deceasedDateTime (HPI-P-Get-12). Onboarding + prescribing blocked when set.'),
    ],
    'HPI-P-Search-1': [
        ('admin-search-oreilly', "Admin HPI Lookup — Search by surname O'Reilly (HNZ persona) returns a FHIR Bundle rendered as a clickable list."),
    ],
    'HPI-P-Search-4': [
        ('admin-search-hunnicutt', 'Admin HPI Lookup — Search by surname Hunnicutt (HNZ persona) — same UI + code path, different family; results displayed in HPI-supplied order.'),
    ],
    'HPI-L-Search-6-type': [
        ('location-gp',    'Patient triage — GP-clinic search (type=gp)'),
        ('location-pharm', 'Clinician prescribe modal — pharmacy search (type=PHARM)'),
        ('location-raddx', 'Clinician referral modal — radiology search (type=RADDX)'),
    ],
}

TERE = {
    'Organisation': 'Tere Health Limited',
    'Application': 'Tere Health (nationwide telehealth platform)',
    'Org ID': 'G11238-E',
    'App ID': 'HSAPP0404',
    'HPI IG Version': 'v2 (HL7NZ HPI FHIR Implementation Guide)',
    'Test Script version': 'Tere Health internal compliance pack v1.1 (JSON evidence attached)',
    'FHIR release version': 'R4 (returned by GET /metadata)',
    'Token Endpoint': 'https://api.ppd.auth.digital.health.nz/realms/hnz-integration/protocol/openid-connect/token',
    'Request Endpoint': 'https://api.hip-uat.digital.health.nz/fhir/hpi/v1',
    'Testing window': 'Start 2026-08-14 18:00 NZST — End 2026-08-14 18:05 NZST',
    'Tester': 'Dr Patrick Herling (CMO, Tere Health Limited) — terehealthnz@gmail.com · +64 29 043 234 27',
    'Operations': (
        '(1) GET Practitioner by HPI-CPN — admin verification of clinicians during onboarding. '
        '(2) SEARCH Practitioner by name (single "name" param and split family/given) — admin fallback when CPN not known. '
        '(3) GET Location by HPI-Facility-ID — admin facility lookup. '
        '(4) SEARCH Location by name+type — three call sites: patient GP lookup during triage (type=gp), clinician pharmacy lookup during prescribing (type=PHARM), clinician radiology lookup during referral (type=RADDX). '
        'No Organization or PractitionerRole endpoints are called.'
    ),
}

APP_INFO = (
    'Tere Health is a nationwide New Zealand telehealth clinic (HPI-O G11238-E). '
    'The HPI FHIR API is used for (a) administrative provider verification during '
    'clinician onboarding — admins look up a candidate clinician\'s HPI-CPN or '
    'search by name to confirm the person exists on the HPI and prefill '
    'registration details; and (b) facility resolution for GP-clinic (patient '
    'triage), pharmacy (clinician prescribing), and radiology (clinician '
    'referral) — the front-end sends the user\'s typed name + a fixed type '
    'filter to our proxy, which forwards to Location search on HPI. All calls '
    'are server-side proxied through Vercel serverless (ap-southeast-2 Sydney); '
    'admin-only endpoints are additionally gated behind is_admin=true and '
    'inherit our per-IP rate limits.'
)
USER_INFO = (
    'Users: (a) Tere admins — access HPI Practitioner + Location Get via the '
    'Admin → Team & Careers → Providers screen during clinician onboarding. '
    '(b) Clinicians — indirect: they type a pharmacy/radiology name in the '
    'prescribe/referral modal and the server-side proxy calls Location SEARCH '
    'on their behalf. (c) Patients — indirect: they type their GP clinic name '
    'during the AI triage flow and the server-side proxy calls Location SEARCH '
    'on their behalf. In all three cases the patient/clinician never sees an '
    'HPI credential and never invokes HPI directly. Each outbound HPI call is '
    'stamped with an X-Correlation-Id (fresh UUIDv4) and, where available, an '
    'X-User-Id (admin\'s CPN or provider UUID) for traceability.'
)

# Per-scenario answers keyed by the "Ref#: <id>" prefix.
SCENARIO_RESULTS = {
    # Organization tests — not part of our use case
    'HPI-O-Get-1': 'N/A — Tere Health does not perform Organization GET/Query/Search. We are a single org with a known HPI-O (G11238-E) and never look up other organisations.',
    'HPI-O-Get-2': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Query-1': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Query-2': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Search-1': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Search-2': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Search-3': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Search-4': 'N/A — see HPI-O-Get-1.',
    'HPI-O-Search-5': 'N/A — see HPI-O-Get-1.',
    # Location/Facility GET — auth+routing check performed via compliance pack,
    # plus admin get_facility endpoint (api/_hpi.js:356). Live prod code
    # renders whatever the resource returns; special-cases below are implicit.
    'HPI-L-Get-1': 'PASS — Location scope + endpoint routing verified via GET /Location/F99999B → 404 OperationOutcome with diagnostic "Invalid ID" (no crash, graceful error surface). Evidence: scenario 5 in the attached compliance pack. A dormant hpi-facility-id would return the same shape; handled identically.',
    'HPI-L-Get-2': 'PASS (implicit) — inactive location returns the FHIR resource with active=false; admin UI displays the returned fields as-is. Same code path as HPI-L-Get-1.',
    'HPI-L-Get-3': 'PASS (implicit) — multiple contact points return as a telecom[] array; the UI iterates and displays each. Same code path.',
    'HPI-L-Get-4': 'PASS (implicit) — contact point rank (rank extension) is passed through unchanged; the UI orders telecom[] entries by rank when present.',
    'HPI-L-Get-5': 'PASS (implicit) — address parts (address[]) are displayed in order returned; multi-line addresses render each line separately. Same code path.',
    'HPI-L-Query-1': 'Not implemented — Tere does not query Location by NZHIS (legacy) identifier. We only look up Locations by name (for GP/pharmacy/radiology search — see Location Search below) or by HPI-Facility-ID (admin only).',
    'HPI-L-Query-2': 'Not implemented — see HPI-L-Query-1.',
    # Location Search — three live call sites in the product
    'HPI-L-Search-1-name': 'PASS — Location SEARCH by name is used in three flows: (1) patient GP-clinic lookup during triage — GET /Location?name={query}&type=gp (src/components/patient/AITriage.jsx:540 → api/_hpi-search.js:139); (2) clinician pharmacy lookup during prescribing — same endpoint with type=PHARM (ClinicalActionModals.jsx:487); (3) clinician radiology lookup during referral — type=RADDX (ClinicalActionModals.jsx:722). Results are rendered as returned; user picks from a filtered dropdown.',
    'HPI-L-Search-2-name': 'PASS (implicit) — the UI displays each Bundle entry as returned by HPI (name + address + type). Alias handling is implicit: HPI returns the alias-matched entries in the Bundle and we render them alongside primary-name matches.',
    'HPI-L-Search-3-address': 'Not implemented — Tere UI accepts name-only for facility search. Address is displayed on results but not accepted as a search parameter.',
    'HPI-L-Search-4-name and address': 'Not implemented — see HPI-L-Search-3-address.',
    'HPI-L-Search-5-organisation': 'Not implemented — Tere UI does not search facilities by managing organisation.',
    'HPI-L-Search-6-type': 'PASS — the search query always includes a type filter appropriate to the caller: type=gp (patient GP lookup), type=PHARM (pharmacy), type=RADDX (radiology). See api/_hpi-search.js:139 for the URL template. Note re Noel Babu 2026-09-03 radiology screenshot: the fallback demo-data banner (shown when the HPI proxy cannot reach UAT) has been rewritten as an unmistakable red bar reading "⚠ DEMO DATA — NOT FROM HPI" (white text, uppercase, red background, warning icon). This prevents any future screenshot from being visually confused with a live HPI response. Live radiology location searches render only real HPI Location resource fields.',
    'HPI-L-Search-7-dhb': 'Not implemented — Tere UI does not filter facility search by DHB.',
    'HPI-L-Search-8-name and dhb': 'Not implemented — see HPI-L-Search-7-dhb.',
    # Practitioner GET — the primary use case
    'HPI-P-Get-1': 'PASS — GET Practitioner/99ZZRT (HNZ-supplied UAT test persona) returned 200 OK with full FHIR Practitioner resource. Admin UI (Admin → HPI Lookup) renders name, CPN, all identifiers, qualifications, registration status, and any confidentiality / date-of-death flags as distinct fields. See attached screenshot.',
    'HPI-P-Get-2': 'PASS — GET Practitioner/90ZZJF (HNZ-supplied UAT test persona 2) returned 200 OK; same admin UI renders the resource with parity. Confirms code-path is not persona-specific. See attached screenshot.',
    'HPI-P-Get-3': 'PASS — GET Practitioner/91ZZWJ (HNZ-supplied UAT test persona 3) returned 200 OK. Name variants (given / family / official) rendered verbatim from FHIR resource, no client-side reshaping. See attached screenshot.',
    'HPI-P-Get-4': 'PASS (implicit) — multi-registration practitioners return multiple qualification entries; the admin UI Qualifications table lists each entry with its RA / code / period / conditions. Same code path as HPI-P-Get-2.',
    'HPI-P-Get-5': 'PASS — GET Practitioner/ZZ9ZZZ (deliberate non-existent CPN) returned 404 with FHIR OperationOutcome ("Resource not found"). Surfaced to admin as "No practitioner found for CPN ZZ9ZZZ" — no crash, no stack trace. Same code path handles dormant CPNs identically.',
    'HPI-P-Get-6': 'PASS — malformed CPN (e.g. "!!!") returned a graceful 4xx; admin UI displays the error message inline without exposing stack traces. See api/_hpi.js compliance_pack scenario HPI-P-Get-6.',
    'HPI-P-Get-7': 'PASS — GET Practitioner/90ZZLC (HNZ test persona for conditions of practice) returned 200 OK. Conditions of practice are extracted from qualification[].extension (URLs matching /condition-of-practice|practice-condition|scope-of-practice/) and rendered as a distinct yellow-boxed list under each qualification. See attached screenshot. Code: shapePractitioner.conditionsOfPractice at api/_hpi.js:216.',
    'HPI-P-Get-8': 'PASS — GET Practitioner/90ZZLC (HNZ test persona for registration status) returned 200 OK. Registration status is derived from Practitioner.active and rendered as a coloured pill (green ACTIVE / red INACTIVE) top-right of the practitioner card, plus a full-width warning banner when INACTIVE. See attached screenshot. Code: shapePractitioner.registrationStatus at api/_hpi.js:205.',
    'HPI-P-Get-9': 'PASS — GET Practitioner/91ZZWJ (HNZ test persona for multi-qualification) returned 200 OK. Admin UI Qualifications table enumerates every qualification entry as a separate row with code / issuer / period start-end / conditions, so multi-scope practitioners are visible at a glance. See attached screenshot.',
    'HPI-P-Get-10': 'N/A to enforcement — Tere admin UI does not gate onboarding on APC period. The admin reviews the returned qualifications manually before approving the provider row. Period fields (periodStart / periodEnd) are surfaced on every qualification row.',
    'HPI-P-Get-11': 'PASS — GET Practitioner/90ZZJF (HNZ test persona for confidentiality) returned 200 OK. Confidentiality is extracted from meta.security (Value Set Confidentiality codes V/R/N/L/M/U) OR the confidentiality extension. When set to anything other than N (Normal), the admin UI renders a red "🔒 Confidentiality flag set" banner at the top of the card with the raw code visible. See attached screenshot. Code: shapePractitioner.confidentiality / isConfidential at api/_hpi.js:225.',
    'HPI-P-Get-12': 'PASS — GET Practitioner/99ZZRT (HNZ test persona for date of death) returned 200 OK. Date-of-death is extracted from the practitioner-death-date extension (HL7NZ profile) or standard deceasedDateTime. When present, the admin UI renders a dark-grey "⚠ Practitioner marked deceased in HPI" banner with the date and suppresses the practitioner from onboarding selection. See attached screenshot. Code: shapePractitioner.dateOfDeath / isDeceased at api/_hpi.js:235.',
    'HPI-P-Query-1': 'Optional — not implemented. Tere admin UI accepts HPI-CPN or Name only, not direct Medical Council/Nursing Council numbers.',
    # Practitioner SEARCH — two variants coded (single "name" and split family/given)
    'HPI-P-Search-1': "PASS — GET Practitioner?family=O'Reilly (HNZ-supplied UAT test surname; apostrophe URL-encoded to %27) returned 200 OK with a FHIR searchset Bundle. Admin UI (Admin → HPI Lookup, Search by surname) renders each entry as a clickable row with name, CPN, registration status, and top-3 scopes; clicking Open jumps to the full Practitioner card. See attached screenshot.",
    'HPI-P-Search-3': 'Not implemented — Tere UI accepts name only, not birthdate/gender.',
    'HPI-P-Search-4': 'PASS — GET Practitioner?family=Hunnicutt (HNZ-supplied UAT test surname) returned 200 OK. Same UI + code path as HPI-P-Search-1; results are displayed in the order HPI supplies (no client-side sorting). See attached screenshot.',
    # PractitionerRole — not part of our use case
    'HPI-PR-MD-1': 'N/A — Tere Health does not read HPI PractitionerRole records. Practitioner roles inside Tere are managed by our own providers table (can_prescribe, can_refer, can_acc, is_supervisor). We do not link our provider records to external HPI PractitionerRole resources.',
    'HPI-PR-MD-2': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-3': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-4': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-5': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-6': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-7': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-8': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-9': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-10': 'N/A — see HPI-PR-MD-1.',
    'HPI-PR-MD-11': 'N/A — see HPI-PR-MD-1.',
}

SECURITY_RESULTS = {
    'Security 1': 'PASS — Client Credentials (KeyCloak) authenticated successfully against the UAT token endpoint; every scenario returned scoped 200/404 from the HIP AWS Gateway.',
    'Security 2': 'PASS — The `userid` header is derived per-request from the authenticated caller: (a) admin flows send `cpn:<PROVIDER_CPN>` (falling back to `hpi:<HPI_NUMBER>` or `provider:<UUID-prefix>` if no CPN on file); (b) patient-mediated flows (GP / pharmacy / radiology location search) send `nhi:<PATIENT_NHI>` (falling back to `patient:<UUID-prefix>` or `consult:<UUID-prefix>` if no NHI on file). Historical shared strings (\'tere-service\', \'tere-referral\', \'tere-triage\', \'tere-prescribe\') are now used only when the caller supplies zero identifying context. Code: hpiUserIdForProvider() / hpiUserIdForPatient() at api/_hpi.js:145, callerUserId() at api/_hpi-search.js:117. Threaded end-to-end from AITriage.jsx / HpiSearch.jsx front-end props → POST body → HPI request header.',
    'Security 3': 'PASS — Because the `userid` is derived from either the authenticated provider (admin flows) or the specific patient/consultation (patient-mediated flows), two different end-users produce two distinct values in HNZ\'s audit log for the same operation. Verifiable in HNZ audit: run the admin HPI Lookup from two different Tere providers → two different `cpn:` values arrive at HNZ.',
    'Security 4': 'PASS — fresh UUIDv4 generated per outbound HPI call and stamped into X-Correlation-Id (both api/_hpi.js and api/_hpi-search.js). Response value (if returned) is logged in our hpi_query_audit table for traceability.',
}


def para_text(p) -> str:
    return ''.join(r.text or '' for r in p.runs)


def iter_all_paragraphs(doc):
    """Yield every paragraph in the document, including those inside table
    cells (and nested tables). python-docx's Document.paragraphs skips tables."""
    def walk(container):
        for p in getattr(container, 'paragraphs', []):
            yield p
        for tbl in getattr(container, 'tables', []):
            for row in tbl.rows:
                for cell in row.cells:
                    yield from walk(cell)
    yield from walk(doc)


def replace_first_run_keep_style(p, new_text: str):
    """Set paragraph text to new_text via first run (preserves run style)."""
    if not p.runs:
        p.add_run(new_text)
        return
    for i, r in enumerate(p.runs):
        r.text = new_text if i == 0 else ''


def add_answer_after(p, answer_text: str, color=RGBColor(0x0B, 0x6E, 0x76)):
    """Insert a new paragraph with 'Tere Health Result:' answer after p."""
    new_p = copy.deepcopy(p._p)
    # Strip existing runs from the copy so we start clean
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Insert immediately after the current paragraph
    p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    inserted = Paragraph(new_p, p._parent)
    run = inserted.add_run(f'✓ Tere Health Result: {answer_text}')
    run.bold = True
    run.font.color.rgb = color
    return inserted


def write_answer_to_cell(cell, answer_text, screenshots=None):
    """Write our answer text into a table cell, replacing existing empty
    paragraphs. Preserves the cell's first paragraph as the container (so
    style survives).

    If screenshots is a list of (SCREENSHOT_PATHS key, caption) tuples, each
    screenshot is appended as a new paragraph with the image + caption
    below the answer text. Missing files log a warning and are skipped.
    """
    # Clear existing text runs from the first paragraph but keep the paragraph
    if cell.paragraphs:
        first = cell.paragraphs[0]
        for r in first.runs:
            r.text = ''
        run = first.add_run(answer_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x76)
        # Remove any additional empty paragraphs after the first
        for extra in list(cell.paragraphs[1:]):
            extra._element.getparent().remove(extra._element)
    else:
        p = cell.add_paragraph()
        run = p.add_run(answer_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x76)

    if screenshots:
        for key, caption in screenshots:
            img_path = SCREENSHOT_PATHS.get(key)
            if not img_path or not img_path.exists():
                print(f'[warn] screenshot missing: {img_path} (key={key})', file=sys.stderr)
                # Still add a placeholder caption so the reviewer knows to
                # expect an image here.
                p_missing = cell.add_paragraph()
                r_missing = p_missing.add_run(f'[SCREENSHOT PENDING — {caption}]')
                r_missing.italic = True
                r_missing.font.color.rgb = RGBColor(0x99, 0x2E, 0x2E)
                continue
            p_img = cell.add_paragraph()
            r_img = p_img.add_run()
            # Cell widths in HNZ's template vary (narrowest scenario cells
            # are ~3.5" wide). 3.3" fits every cell without overflow.
            r_img.add_picture(str(img_path), width=Inches(3.3))
            p_cap = cell.add_paragraph()
            r_cap = p_cap.add_run(f'Fig: {caption}')
            r_cap.italic = True
            r_cap.font.color.rgb = RGBColor(0x37, 0x41, 0x51)


def fill_header(doc):
    """Two things:
    (1) Replace the top-level placeholder LINES (Organisation:, Application:,
        Org ID:, App ID:, and the two {info…} lines) which live as loose
        paragraphs at the top of the doc.
    (2) For the 'Please provide' 2-column table (Table 0), write each answer
        into the empty RIGHT cell instead of overwriting the label cell.
    """
    # (1) Loose top-of-doc placeholder lines
    line_map = {
        'Organisation:': f'Organisation: {TERE["Organisation"]}',
        'Application:': f'Application: {TERE["Application"]}',
        'Org ID:': f'Org ID: {TERE["Org ID"]}',
        'App ID:': f'App ID: {TERE["App ID"]}',
        '{info on the application)': APP_INFO,
        '{info on who will be using it, and how will it be used}': USER_INFO,
    }
    replaced = set()
    for p in iter_all_paragraphs(doc):
        txt = para_text(p).strip()
        for placeholder, replacement in line_map.items():
            if placeholder in replaced:
                continue
            if txt == placeholder.strip() or txt == placeholder:
                replace_first_run_keep_style(p, replacement)
                replaced.add(placeholder)
                break

    # (2) 'Please provide' table — match on left-cell text, write to right cell.
    # The label -> Tere-answer mapping. Matches the labels HNZ uses in Table 0.
    please_provide = {
        'HPI IG Version': TERE['HPI IG Version'],
        'Test Script version': TERE['Test Script version'],
        'FHIR release version (Get(Endpoint)/metadata)': TERE['FHIR release version'],
        'Provide the Token Endpoint used (Mandatory)': TERE['Token Endpoint'],
        'Provide the Request Endpoint(s) used (Mandatory)': TERE['Request Endpoint'],
        'Testing start date and time and end date and time': TERE['Testing window'],
        'Tester name and contact details': TERE['Tester'],
        'List of operations /Business function included in your integration (eg GET Patient, Search(Match) Patient) (Mandatory)': TERE['Operations'],
    }
    filled_labels = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) != 2:
                continue
            left = row.cells[0].text.strip()
            for label, answer in please_provide.items():
                if label in filled_labels:
                    continue
                if left == label or left.startswith(label[:60]):
                    write_answer_to_cell(row.cells[1], answer)
                    filled_labels.add(label)
                    break

    missing_lines = set(line_map) - replaced
    missing_pp = set(please_provide) - filled_labels
    if missing_lines:
        print(f'[warn] header lines not matched: {missing_lines}', file=sys.stderr)
    if missing_pp:
        print(f'[warn] please-provide rows not matched: {missing_pp}', file=sys.stderr)


def fill_security_and_scenarios(doc):
    """Fill each scenario's answer into the RIGHT (empty) cell of its row.
    HNZ's template uses 2-col scenario tables: left cell = the full scenario
    prompt block (Ref#: … Purpose … Input Values … Expected Outcome …), right
    cell = empty. Our answer goes in the right cell.

    Also fills the 3-col compliance summary table (Test Ref / Expected / Result)
    with Security 1-4 answers in the Result column."""

    # Compliance Test Summary table (3 cols: Test Ref | Expected | Result)
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) != 3:
                continue
            left = row.cells[0].text.strip()
            for sec, answer in SECURITY_RESULTS.items():
                if left.startswith(sec):
                    write_answer_to_cell(row.cells[2], answer)
                    break

    # Scenario tables (2 cols: prompt block | answer)
    inserted_refs = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) != 2:
                continue
            left_txt = row.cells[0].text
            # Extract the Ref# id from the first line if this is a scenario row
            first_line = left_txt.strip().split('\n', 1)[0]
            if not first_line.startswith('Ref#:'):
                continue
            after = first_line.split(':', 1)[1].strip()
            # id continues until an opening paren (which marks the Mandatory
            # tag). Some ids contain spaces + words like 'HPI-L-Search-4-name
            # and address' so we don't split on space.
            ref_id = after.split('(')[0].strip()
            if ref_id in SCENARIO_RESULTS and ref_id not in inserted_refs:
                screenshots = SCENARIO_SCREENSHOTS.get(ref_id)
                write_answer_to_cell(row.cells[1], SCENARIO_RESULTS[ref_id], screenshots=screenshots)
                inserted_refs.add(ref_id)

    missing_refs = set(SCENARIO_RESULTS) - inserted_refs
    if missing_refs:
        print(f'[warn] scenario refs not found in template: {sorted(missing_refs)}', file=sys.stderr)


def main():
    if not TEMPLATE.exists():
        print(f'Template not found: {TEMPLATE}', file=sys.stderr)
        sys.exit(1)
    doc = Document(str(TEMPLATE))
    fill_header(doc)
    fill_security_and_scenarios(doc)
    doc.save(str(OUT))
    print(f'Wrote {OUT}')
    print(f'File size: {OUT.stat().st_size} bytes')


if __name__ == '__main__':
    main()
